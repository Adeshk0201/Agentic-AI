import os
from fastapi import FastAPI, UploadFile, File, Request
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document

# Import your specialist agents and coordinator logic
from Utils.Agents import (
    Cardiologist, Psychologist, Pulmonologist,
    Neurologist, Endocrinologist, Nutritionist, Coordinator
)

from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()

# Add CORS middleware
origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files and set up templates
# app.mount("/", StaticFiles(directory="frontend/build", html=True), name="frontend")
templates = Jinja2Templates(directory="templates")

# Homepage
@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

# Diagnose and generate .docx report
@app.post("/diagnose/")
async def diagnose(file: UploadFile = File(...)):
    uploads_dir = "uploads"
    os.makedirs(uploads_dir, exist_ok=True)
    file_location = os.path.join(uploads_dir, file.filename)

    # Save uploaded file
    with open(file_location, "wb") as f:
        f.write(await file.read())

    # Read uploaded text
    with open(file_location, "r", encoding="utf-8") as f:
        medical_report = f.read()

    print("Medical report content:")
    print(medical_report)

    # Run through coordinator for diagnosis
    coordinator = Coordinator(medical_report)
    results = coordinator.run_full_evaluation()

    print("Diagnosis results:")
    print(results)

    final_diagnosis = results["multidisciplinary_team"]

    # Save diagnosis as .docx
    results_dir = "results"
    os.makedirs(results_dir, exist_ok=True)
    doc = Document()
    doc.add_heading("Final Diagnosis (Refined)", level=1)
    doc.add_paragraph(final_diagnosis)
    docx_path = os.path.join(results_dir, "final_diagnosis.docx")
    doc.save(docx_path)

    # Return result to frontend
    return JSONResponse(content={"result": final_diagnosis})

# Download endpoint for .docx report
@app.get("/download/", response_class=FileResponse)
async def download():
    doc_path = "results/final_diagnosis.docx"
    if os.path.exists(doc_path):
        return FileResponse(
            path=doc_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="final_diagnosis.docx"
        )
    else:
        return JSONResponse(status_code=404, content={"error": "File not found"})

# Run the app
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8005)


